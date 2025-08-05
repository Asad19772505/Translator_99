import streamlit as st
import os
import tempfile
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF

# Initialize GROQ client using Streamlit secrets
client = Groq(api_key=st.secrets["GROQ_API_KEY"])

def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def translate_text_with_groq(text):
    prompt = f"Translate the following text into English:\n{text}"
    response = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

def save_to_word(translated_text):
    doc = Document()
    doc.add_heading('Translated Document', 0)
    doc.add_paragraph(translated_text)
    temp_word = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_word.name)
    return temp_word.name

def save_to_pdf(translated_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in translated_text.split('\n'):
        pdf.multi_cell(0, 10, line)
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_pdf.name)
    return temp_pdf.name

# Streamlit UI
st.title("PDF Translator using GROQ AI")
st.markdown("Upload a PDF in any language and get the translated document in English.")

uploaded_file = st.file_uploader("Upload a PDF", type=["pdf"])
output_format = st.selectbox("Select output format", ["PDF", "Word"])

if uploaded_file and st.button("Translate"):
    with st.spinner("Extracting text from PDF..."):
        pdf_text = extract_text_from_pdf(uploaded_file)

    if pdf_text.strip() == "":
        st.error("No text could be extracted from the uploaded PDF.")
    else:
        with st.spinner("Translating text using GROQ AI..."):
            translated = translate_text_with_groq(pdf_text)

        with st.spinner(f"Saving translated text as {output_format}..."):
            if output_format == "PDF":
                output_file = save_to_pdf(translated)
                st.success("Translation complete. Download your PDF below.")
                with open(output_file, "rb") as f:
                    st.download_button("Download PDF", f, file_name="translated_document.pdf")
            else:
                output_file = save_to_word(translated)
                st.success("Translation complete. Download your Word file below.")
                with open(output_file, "rb") as f:
                    st.download_button("Download Word Document", f, file_name="translated_document.docx")
